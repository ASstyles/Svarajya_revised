from __future__ import annotations

import csv
import hashlib
import re
from collections import Counter, defaultdict
from difflib import SequenceMatcher
from pathlib import Path


BASE = Path("/Users/amitkothari/Documents/New project/Svarajya-Codex-Audit")
OUT = BASE / "audit-output/04-bugs-and-qa"
INPUT_BUG = BASE / "audit-input/03-bug-and-evidence"
INPUT_TEAM = BASE / "audit-input/02-existing-team-outputs"
INPUT_EXTRA = BASE / "audit-input/06-additional-qa-ux-stress-evidence"
ARCH = BASE / "audit-output/01-repository-and-architecture"
DB = BASE / "audit-output/02-database-and-dependencies"
PRD = BASE / "audit-output/03-prd-traceability"
HANDOVER = BASE / "audit-output/06-final-handover"


MASTER_HEADERS = [
    "Master Bug ID",
    "Source Bug ID",
    "Source File / Sheet",
    "Source Row / Evidence ID",
    "Module",
    "Screen / Route",
    "Bug Title",
    "Bug Description",
    "Bug Type",
    "Severity",
    "Priority Reason",
    "Reproduction Status",
    "Reproduction Steps",
    "Expected Result",
    "Actual Result",
    "Evidence Available",
    "Evidence Link / Evidence ID",
    "Likely Page / Component File",
    "Likely Hook / Store / Service File",
    "Likely API Route",
    "Likely Supabase Table",
    "Likely DB Field",
    "Related Modules",
    "Dashboard Impact",
    "Reminder Impact",
    "Document Vault Impact",
    "Auth / Security Impact",
    "Data Loss Risk",
    "Calculation Risk",
    "Duplicate Status",
    "Duplicate Of",
    "Technical Confidence",
    "AI Confidence Reason",
    "Needs WD Verification",
    "Suggested Jira Issue Type",
    "Suggested Jira Title",
    "Recommended Next Action",
    "Assigned To",
    "Status",
    "Notes",
]

RAW_HEADERS = [
    "Raw Source Row ID",
    "Source Name",
    "Source File / Sheet",
    "Owner Team",
    "Source Row / Evidence ID",
    "Source Bug ID",
    "Module",
    "Screen / Route",
    "Title",
    "Description / Finding",
    "Raw Severity / Priority",
    "Raw Status",
    "Evidence Reference",
    "Triage Outcome",
    "Triage Reason",
    "Include In Master",
    "Master Candidate ID",
    "Duplicate Candidate",
]

SOURCE_HEADERS = [
    "Source Name",
    "File Name",
    "Owner Team",
    "Date",
    "Type of Bugs",
    "Evidence Available",
    "Rows Reviewed",
    "Rows Sent To Raw Merge",
    "Rows Included In Master",
    "Notes",
]

BUG_TYPES = {
    "Auth",
    "Save/Create",
    "Edit/Update",
    "Delete",
    "Data Persistence",
    "Dashboard Calculation",
    "Reminder",
    "Document Upload",
    "Navigation",
    "Validation",
    "Responsive/UI",
    "Duplicate Record",
    "Data Mismatch",
    "Performance",
    "Configuration",
    "Unknown",
}
REPRO_STATUSES = {
    "Reproduced Twice",
    "Reproduced Once",
    "Not Reproduced",
    "Cannot Reproduce",
    "Needs Dummy Login",
    "Needs WD Verification",
    "Code Risk Only",
}
DUP_STATUSES = {
    "Unique",
    "Duplicate",
    "Possible Duplicate",
    "Parent Bug",
    "Merged Into Parent",
    "Needs Review",
}
CONFIDENCE = {"Confirmed", "Strong Evidence", "Probable", "Unconfirmed"}
ACTIONS = {
    "Create Jira Bug",
    "Merge with Existing Jira",
    "Needs WD Reproduction",
    "Needs PM Retest",
    "Needs DA Evidence",
    "Needs Himanshu Review",
    "Needs Harsh Review",
    "Needs Supabase Verification",
    "Needs Auth Review",
    "Move to P0/P1 Impact Report",
    "No Action — Duplicate",
    "No Action — Cannot Verify",
}


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, headers: list[str], rows: list[dict[str, str]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=headers, extrasaction="ignore")
        w.writeheader()
        for row in rows:
            w.writerow({h: clean(row.get(h, "")) for h in headers})


def clean(value) -> str:
    s = "" if value is None else str(value)
    s = s.replace("\x00", " ")
    s = re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "dummyuser@example.com", s)
    s = re.sub(r"(access_token=)[^&\s]+", r"\1[redacted]", s)
    s = re.sub(r"(refresh_token=)[^&\s]+", r"\1[redacted]", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def simple_path(rel: Path) -> str:
    try:
        return str(rel.relative_to(BASE))
    except ValueError:
        return str(rel)


def sha12(path: Path) -> str:
    if not path.exists():
        return ""
    return hashlib.sha256(path.read_bytes()).hexdigest()[:12]


def norm_tokens(text: str) -> set[str]:
    words = re.findall(r"[a-z0-9]+", text.lower())
    stop = {"the", "and", "with", "for", "from", "into", "field", "fields", "issue", "bug", "page", "form", "screen"}
    return {w for w in words if len(w) > 2 and w not in stop}


def similarity(a: str, b: str) -> float:
    ta, tb = norm_tokens(a), norm_tokens(b)
    if not ta or not tb:
        return 0.0
    jaccard = len(ta & tb) / len(ta | tb)
    seq = SequenceMatcher(None, a.lower(), b.lower()).ratio()
    return max(jaccard, seq * 0.75)


def canonical_module(text: str) -> str:
    t = text.lower()
    checks = [
        ("Authentication", ["auth", "login", "oauth", "google", "password reset", "otp", "session", "signup", "sign-up"]),
        ("Onboarding", ["onboarding", "dob", "date of birth", "verification", "mobile otp"]),
        ("Sthapana / Foundation", ["sthapana", "foundation", "family", "profile", "education", "qualification"]),
        ("Pehchaan / Identity Vault", ["pehchaan", "identity", "kyc", "pan", "aadhaar"]),
        ("Kunji / Credentials", ["kunji", "dwaar", "credential", "password vault", "portal"]),
        ("Kosh / Income", ["kosh", "income", "deduction", "gross income", "salary"]),
        ("Vyaya / Expenses", ["vyaya", "expense", "subscription", "leakage", "recurring"]),
        ("Pravah / Banking", ["pravah", "bank", "cash", "liquidity", "account"]),
        ("Rin / Loans", ["rin", "loan", "debt", "emi", "liability"]),
        ("Raksha / Insurance", ["raksha", "insurance", "policy", "premium", "sum assured"]),
        ("Nivesh / Investments", ["nivesh", "investment", "investments", "roi"]),
        ("Bhoomi / Property", ["bhoomi", "property", "real estate"]),
        ("Kar / Tax", ["kar", "tax", "itr", "gst", "din", "financial year", "assessment year"]),
        ("Mitra / Legacy", ["mitra", "will", "nominee", "legacy", "heir"]),
        ("Uttaradhikar / Succession", ["uttaradhikar", "succession", "executor", "emergency contact", "emergency"]),
        ("Rajya / Dashboard", ["rajya", "dashboard", "score", "map", "mandala"]),
        ("Doot / Reminders", ["doot", "reminder", "notification", "alerts"]),
        ("Document Vault", ["document", "vault", "nidhi", "granthagaar", "drive", "upload", "file"]),
        ("Performance / Stress", ["stress", "performance", "k6", "load", "slow", "latency"]),
        ("General UX", ["ux", "ui", "responsive", "button", "touch", "layout", "navigation", "back button"]),
    ]
    for name, keys in checks:
        if any(k in t for k in keys):
            return name
    return clean(text) or "Unknown"


def module_prefix(module: str) -> str:
    mapping = {
        "Authentication": "AUTH",
        "Onboarding": "ONBOARD",
        "Sthapana / Foundation": "PROFILE",
        "Pehchaan / Identity Vault": "IDENTITY",
        "Kunji / Credentials": "CREDENTIAL",
        "Kosh / Income": "INCOME",
        "Vyaya / Expenses": "EXPENSE",
        "Pravah / Banking": "BANK",
        "Rin / Loans": "LOAN",
        "Raksha / Insurance": "INSURANCE",
        "Nivesh / Investments": "INVEST",
        "Bhoomi / Property": "PROPERTY",
        "Kar / Tax": "TAX",
        "Mitra / Legacy": "LEGACY",
        "Uttaradhikar / Succession": "SUCCESSION",
        "Rajya / Dashboard": "DASH",
        "Doot / Reminders": "REMINDER",
        "Document Vault": "DOC",
        "Performance / Stress": "PERF",
        "General UX": "UX",
    }
    return mapping.get(module, "GEN")


def infer_type(text: str) -> str:
    t = text.lower()
    if any(k in t for k in ["oauth", "login", "password", "otp", "session", "auth", "signup", "sign-up"]):
        return "Auth"
    if any(k in t for k in ["save", "create", "add ", "submit"]):
        return "Save/Create"
    if any(k in t for k in ["edit", "update", "change", "modify"]):
        return "Edit/Update"
    if any(k in t for k in ["delete", "remove", "unlink"]):
        return "Delete"
    if any(k in t for k in ["persist", "lost", "not stored", "not saved", "localstorage", "indexeddb", "null"]):
        return "Data Persistence"
    if any(k in t for k in ["dashboard", "score", "calculation", "calculate", "gross", "decimal", "amount", "emi", "roi"]):
        return "Dashboard Calculation"
    if any(k in t for k in ["reminder", "notification", "alert", "due date", "renewal"]):
        return "Reminder"
    if any(k in t for k in ["upload", "document", "file", "vault", "drive", "opfs"]):
        return "Document Upload"
    if any(k in t for k in ["route", "redirect", "back button", "navigation", "404", "callback"]):
        return "Navigation"
    if any(k in t for k in ["validation", "invalid", "future", "past", "digit", "letters", "regex", "format", "mandatory"]):
        return "Validation"
    if any(k in t for k in ["ui", "ux", "responsive", "layout", "spacing", "button", "touch", "mobile", "desktop"]):
        return "Responsive/UI"
    if any(k in t for k in ["duplicate", "twice"]):
        return "Duplicate Record"
    if any(k in t for k in ["mismatch", "conflict", "wrong", "unclear", "mapping"]):
        return "Data Mismatch"
    if any(k in t for k in ["performance", "stress", "load", "slow", "latency", "k6"]):
        return "Performance"
    if any(k in t for k in ["config", "vercel", "environment", "supabase", "firebase", "rls", "bucket"]):
        return "Configuration"
    return "Unknown"


def severity(raw: str, title: str, desc: str, btype: str) -> str:
    s = (raw or "").upper()
    if "P0" in s:
        return "P0"
    if "CRITICAL" in s and btype in {"Auth", "Document Upload"}:
        return "P1"
    if "P1" in s or "HIGH" in s:
        return "P1"
    if "P2" in s or "MEDIUM" in s or "WARNING" in s:
        return "P2"
    if "P3" in s or "LOW" in s or "MINOR" in s:
        return "P3"
    t = f"{title} {desc}".lower()
    if any(k in t for k in ["complete app failure", "destructive production data loss"]):
        return "P0"
    if any(k in t for k in ["blocked", "cannot complete", "404", "password reset", "otp", "file-not-found", "schema", "route regression"]):
        return "P1"
    if btype in {"Responsive/UI"}:
        return "P3"
    return "P2"


def priority_reason(sev: str, btype: str, module: str, title: str) -> str:
    if sev == "P0":
        return "Potential login, security, production data loss, or complete app failure risk; requires immediate verification."
    if sev == "P1":
        return "Important user journey or shared module behavior may be blocked or materially degraded."
    if sev == "P2":
        return "Feature can work but may save, display, calculate, or connect data incorrectly."
    return "Primarily UI, wording, layout, or lower-risk usability issue unless WD evidence shows deeper impact."


def repro_status(raw_status: str, source: str, desc: str) -> str:
    t = f"{raw_status} {desc}".lower()
    if "code risk" in t:
        return "Code Risk Only"
    if "reproduced 2" in t or "re-confirmed" in t or ("attempt 1" in t and "attempt 2" in t):
        return "Reproduced Twice"
    if "confirmed" in t and source.startswith(("PM", "NB", "DA")):
        return "Reproduced Once"
    if "pass" == raw_status.strip().lower() or "not reproduced" in t:
        return "Not Reproduced"
    if "cannot test" in t or "blocked" in t:
        return "Cannot Reproduce"
    if "dummy login" in t:
        return "Needs Dummy Login"
    return "Needs WD Verification"


def confidence_for(source: str, repro: str, desc: str) -> str:
    text = f"{source} {desc}".lower()
    if repro in {"Reproduced Twice", "Reproduced Once"} and any(k in text for k in ["code", "table", "api", "route", "source", "confirmed", "evidence"]):
        return "Strong Evidence"
    if repro in {"Reproduced Twice", "Reproduced Once"}:
        return "Strong Evidence"
    if repro == "Code Risk Only":
        return "Probable"
    if "jira" in text or "report" in text:
        return "Unconfirmed"
    return "Probable"


def build_enrichment():
    module_map = read_csv(ARCH / "MODULE_TO_FILE_MAP.csv")
    crud = read_csv(INPUT_TEAM / "2026-07-19_MODULE_CRUD_TRACE.csv")
    fdb = read_csv(DB / "FRONTEND_DATABASE_FIELD_MAP.csv")
    dep = read_csv(DB / "MODULE_DEPENDENCY_MATRIX.csv")
    sup = read_csv(ARCH / "SUPABASE_CODE_USAGE_MAP.csv")
    dbinv = read_csv(DB / "DATABASE_TABLE_INVENTORY.csv")
    by_module = defaultdict(lambda: defaultdict(list))

    def add(module, key, value):
        value = clean(value)
        if value and value not in by_module[module][key]:
            by_module[module][key].append(value)

    for row in module_map:
        mod = canonical_module(" ".join(row.values()))
        add(mod, "route", row.get("Main_Route", ""))
        add(mod, "page", row.get("Main_Page_File", ""))
        add(mod, "page", row.get("Related_Page_Files", ""))
        add(mod, "api", row.get("API_Routes", ""))
        add(mod, "service", row.get("Service_or_Store", ""))
        add(mod, "table", row.get("Data_or_Storage", ""))
        add(mod, "evidence", f"audit-output/01-repository-and-architecture/MODULE_TO_FILE_MAP.csv:{row.get('Module_ID','')}")
    for row in crud:
        mod = canonical_module(row.get("Module Name", ""))
        add(mod, "page", row.get("1. Form Location", ""))
        add(mod, "api", row.get("2. Save Action", ""))
        add(mod, "table", row.get("3. Table Used", ""))
        add(mod, "service", row.get("4. Edit Action", ""))
        add(mod, "service", row.get("5. Delete Action", ""))
        add(mod, "evidence", f"audit-input/02-existing-team-outputs/2026-07-19_MODULE_CRUD_TRACE.csv:{row.get('Module #','')}")
    for row in fdb:
        mod = canonical_module(row.get("Module", ""))
        add(mod, "page", row.get("Component_File", ""))
        add(mod, "service", row.get("Service_API_Store", ""))
        add(mod, "table", row.get("Database_Table", ""))
        add(mod, "field", row.get("Database_Column", ""))
    for row in dep:
        mod = canonical_module(row.get("Module", ""))
        add(mod, "table", row.get("Primary_Tables", ""))
        add(mod, "dashboard", row.get("Dashboard_Impact", ""))
        add(mod, "reminder", row.get("Reminder_Impact", ""))
        add(mod, "vault", row.get("Document_Vault_Impact", ""))
        add(mod, "calc", row.get("Score_or_Calculation_Impact", ""))
        add(mod, "related", row.get("Reads_From_Modules", ""))
        add(mod, "related", row.get("Writes_For_Modules", ""))
    for row in sup:
        mod = canonical_module(row.get("Module_Association", ""))
        add(mod, "service", row.get("File_Path", ""))
        add(mod, "table", row.get("Resource", ""))
    table_by_name = {r.get("Table", ""): r for r in dbinv}
    return by_module, table_by_name


def route_from(row: dict[str, str], text: str) -> str:
    for key in ["Screen / route", "Screen / Route", "Route / Screen", "Screen/Route"]:
        if row.get(key):
            return row[key]
    m = re.search(r"(/[A-Za-z0-9_\-/]+)", text)
    return m.group(1) if m else "Needs WD Verification"


raw: list[dict[str, str]] = []


def add_raw(source, file_rel, owner, rownum, source_id, module, screen, title, desc, sev="", status="", evidence="", outcome="Valid bug", reason="", include=True):
    module = canonical_module(f"{module} {title} {desc}")
    raw.append({
        "Raw Source Row ID": f"RAW-{len(raw)+1:04d}",
        "Source Name": source,
        "Source File / Sheet": file_rel,
        "Owner Team": owner,
        "Source Row / Evidence ID": str(rownum),
        "Source Bug ID": clean(source_id),
        "Module": module,
        "Screen / Route": clean(screen) or route_from({}, f"{title} {desc}"),
        "Title": clean(title) or "Untitled bug-like finding",
        "Description / Finding": clean(desc),
        "Raw Severity / Priority": clean(sev),
        "Raw Status": clean(status),
        "Evidence Reference": clean(evidence),
        "Triage Outcome": outcome,
        "Triage Reason": clean(reason),
        "Include In Master": "Yes" if include else "No",
        "Master Candidate ID": "",
        "Duplicate Candidate": "",
    })


def gather_sources():
    source_stats = {}

    pm_path = INPUT_BUG / "PM_VERIFIED_BUGS.csv"
    pm = read_csv(pm_path)
    for i, r in enumerate(pm, start=2):
        add_raw("PM verified bug register", simple_path(pm_path), "PM", i, f"PM-{r.get('Bug #','')}", r.get("Module (per report)", ""), "", r.get("Short title", ""), " | ".join([r.get("Attempt 1 result", ""), r.get("Attempt 2 result", ""), r.get("Notes / prior observations (NOT formal verification)", "")]), r.get("Suggested severity") or r.get("Reported severity"), r.get("Verified status", ""), r.get("Duplicate of", ""), "Valid bug", "PM verified register row retained; WD verification still required.", True)
    source_stats["PM verified bug register"] = (pm_path, "PM", len(pm), len(pm), len(pm), "User journey bugs with PM verification notes")

    nb_path = INPUT_BUG / "EXISTING_BUG_REPORT.csv"
    nb = read_csv(nb_path)
    for i, r in enumerate(nb, start=2):
        add_raw("Existing/NB bug report", simple_path(nb_path), "QA/PM", i, r.get("NB-ID", ""), r.get("Module", ""), r.get("Screen / route", ""), r.get("Title", ""), r.get("Description", ""), r.get("Suggested severity", ""), r.get("Status", ""), r.get("Related report bug", ""), "Valid bug", "Existing bug report row retained; duplicate linkage resolved in master.", True)
    source_stats["Existing/NB bug report"] = (nb_path, "QA/PM", len(nb), len(nb), len(nb), "Existing bug report and new bug IDs")

    daj_path = INPUT_BUG / "DA_BUG_VERIFICATION.csv"
    daj = read_csv(daj_path)
    included = 0
    for i, r in enumerate(daj, start=2):
        status = r.get("Status", "")
        include = status.strip().lower() != "pass"
        if include:
            included += 1
        outcome = "Valid bug" if include else "Excluded - Passed journey"
        reason = "Non-pass DA journey triaged as bug/verification item." if include else "DA journey passed; retained in raw merge only."
        add_raw("DA bug verification journeys", simple_path(daj_path), "DA", i, r.get("Bug ref (report # / NB-#)", "") or f"DA-JOURNEY-{r.get('Journey','')}-{r.get('Step #','')}", r.get("Journey name", ""), "", f"{r.get('Journey name','')} step {r.get('Step #','')}: {r.get('Step action (from task doc)','')}", r.get("Actual result", ""), r.get("Status", ""), r.get("Status", ""), r.get("Notes / known issues to watch", ""), outcome, reason, include)
    source_stats["DA bug verification journeys"] = (daj_path, "DA", len(daj), len(daj), included, "Journey test results; only non-pass rows included in master")

    p01_path = INPUT_BUG / "P0_P1_BUG_REPORT_DRAFT.csv"
    p01 = read_csv(p01_path)
    for i, r in enumerate(p01, start=2):
        add_raw("P0/P1 bug report draft", simple_path(p01_path), "QA/AI", i, r.get("bug_id", ""), r.get("module", ""), "", r.get("title", ""), " | ".join([r.get("steps_or_attempts", ""), r.get("notes", "")]), r.get("severity", ""), r.get("status", ""), "", "Duplicate/supporting source", "P0/P1 draft is a severity handoff source; rows are retained and linked to parent master bugs.", True)
    source_stats["P0/P1 bug report draft"] = (p01_path, "QA/AI", len(p01), len(p01), len(p01), "P0/P1 severity handoff source")

    das_path = INPUT_TEAM / "DA_SUSPECTED_MISSING_CONNECTIONS.csv"
    das = read_csv(das_path)
    for i, r in enumerate(das, start=2):
        add_raw("DA suspected missing connections", simple_path(das_path), "DA", i, r.get("Issue ID", ""), f"{r.get('Source Module','')} {r.get('Destination Module','')}", "", f"{r.get('Source Module','')} to {r.get('Destination Module','')}: {r.get('Issue Group','')}", " | ".join([r.get("Information Entered or Changed", ""), r.get("Expected Connection", ""), r.get("Actual Observed Behaviour", ""), r.get("Business/User Impact", "")]), r.get("Priority", ""), r.get("Connection Status", ""), r.get("Evidence Reference", ""), "Valid bug", "Dependency issue retained; connection ownership needs WD confirmation.", True)
    source_stats["DA suspected missing connections"] = (das_path, "DA", len(das), len(das), len(das), "Cross-module missing/partial dependency bugs")

    dau_path = INPUT_TEAM / "DA_UNCLEAR_AND_CONFLICTING_FIELDS.csv"
    dau = read_csv(dau_path)
    for i, r in enumerate(dau, start=2):
        itype = r.get("Issue type", "")
        if "duplicate" in itype.lower():
            outcome = "Duplicate/code risk"
        else:
            outcome = "Code risk"
        add_raw("DA unclear/conflicting fields", simple_path(dau_path), "DA", i, r.get("Issue ID", ""), r.get("Module", ""), "", f"{r.get('Frontend field label','')} mapping unclear", " | ".join([r.get("Issue type", ""), r.get("What is unclear or conflicting", ""), r.get("Question for technical team", ""), f"Possible DB: {r.get('Possible database table(s)','')} {r.get('Possible database field(s)','')}"]), r.get("Priority", ""), "Needs technical confirmation", r.get("Source tab", ""), outcome, "All DA unclear-field findings triaged as code risk unless duplicate-table wording is present.", True)
    source_stats["DA unclear/conflicting fields"] = (dau_path, "DA", len(dau), len(dau), len(dau), "Field mismatch and ambiguous frontend/database mapping risks")

    ev_path = INPUT_BUG / "EVIDENCE_LIBRARY.csv"
    ev = read_csv(ev_path)
    source_stats["DA evidence library"] = (ev_path, "DA/QA", len(ev), len(ev), 0, "Supporting evidence library; linked by NB IDs")

    for jira_name in ["Jira.csv", "Jira (1).csv"]:
        jpath = INPUT_EXTRA / jira_name
        rows = read_csv(jpath)
        inc = 0
        for i, r in enumerate(rows, start=2):
            issue_type = r.get("Issue Type", "")
            include = issue_type.lower() == "bug"
            if include:
                inc += 1
            outcome = "Valid bug" if include else "Excluded - Jira non-bug work item"
            reason = "Jira bug row retained." if include else f"Jira Issue Type is {issue_type}; retained in raw only."
            add_raw(f"Jira export {jira_name}", simple_path(jpath), "Jira/PM", i, r.get("Issue key", ""), r.get("Parent summary", "") or r.get("Summary", ""), "", r.get("Summary", ""), r.get("Description", ""), r.get("Priority", ""), r.get("Status", ""), f"Jira issue {r.get('Issue key','')}; parent {r.get('Parent key','')}", outcome, reason, include)
        source_stats[f"Jira export {jira_name}"] = (jpath, "Jira/PM", len(rows), len(rows), inc, "Jira status/priority/parent export; bug rows included in master")

    gather_extra_docs(source_stats)
    return source_stats


def doc_text(path: Path) -> str:
    if path.suffix.lower() == ".md":
        return path.read_text(encoding="utf-8", errors="ignore")
    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for row in t.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    if path.suffix.lower() == ".pdf":
        from pypdf import PdfReader
        return "\n".join((page.extract_text() or "") for page in PdfReader(str(path)).pages)
    return ""


def extract_doc_issues(path: Path) -> list[tuple[str, str, str]]:
    txt = doc_text(path)
    lines = [clean(x) for x in txt.splitlines() if clean(x)]
    issues: list[tuple[str, str, str]] = []
    last = ""
    for idx, line in enumerate(lines):
        if len(line) > 260:
            line = line[:257] + "..."
        is_heading = (
            re.search(r"\bBUG-\d{3}\b", line, re.I)
            or re.match(r"(Issue\s*#?\d+|[0-9]+\.\s+.+\[(High|Medium|Low|Critical)\])", line, re.I)
            or re.match(r"(UX Issue|UI Issue|Current Issue|The Issue|Issue\s*[–:-])", line, re.I)
        )
        if is_heading:
            desc = " ".join(lines[idx + 1 : idx + 4])[:650]
            title = line
            source_id = re.search(r"(BUG-\d{3}|Issue\s*#?\d+|[0-9]+)", line, re.I)
            issues.append((source_id.group(1) if source_id else f"DOC-{len(issues)+1:03d}", title, desc))
            last = title
    # Keep the document represented even when headings were not structurally extractable.
    if not issues and lines:
        issues.append(("DOC-SUMMARY", lines[0], " ".join(lines[1:5])[:650]))
    # Avoid flooding master with every UX subheading; raw rows still preserve document-level evidence.
    limit = 20 if path.name in {"Bug report.pdf", "Sav-Rajya (1).docx", "Stress Test.pdf", "Svarajya_UX_Audit_Analysis_Samiksha_V1 (1).pdf"} else 12
    return issues[:limit]


def gather_extra_docs(source_stats):
    for path in sorted(INPUT_EXTRA.iterdir()):
        if path.name.startswith("Jira") or path.suffix.lower() not in {".pdf", ".docx", ".md"}:
            continue
        issues = extract_doc_issues(path)
        duplicate_upload = path.name == "bugs_fix_implementation_plan (2).md"
        included = 0
        for n, (sid, title, desc) in enumerate(issues, start=1):
            include = not duplicate_upload
            if include:
                included += 1
            outcome = "Duplicate source file" if duplicate_upload else "Report/fix-plan finding"
            reason = "Byte-for-byte duplicate of bugs_fix_implementation_plan (1).md; preserved in raw only." if duplicate_upload else "Clear issue/fix-plan heading extracted from document; needs reconciliation with current app/source."
            add_raw(f"Additional evidence document - {path.name}", simple_path(path), "QA/UX/WD", n, sid, title, "", title, desc, "", "Needs WD Verification", f"sha256:{sha12(path)}", outcome, reason, include)
        source_stats[f"Additional evidence document - {path.name}"] = (path, "QA/UX/WD", len(issues), len(issues), included, "Additional offline document evidence; extracted headings only")


def make_master():
    enrichment, table_by_name = build_enrichment()
    evidence_by_id = {}
    ev_path = INPUT_BUG / "EVIDENCE_LIBRARY.csv"
    for r in read_csv(ev_path):
        evidence_by_id[r.get("reference_id", "")] = r

    included = [r for r in raw if r["Include In Master"] == "Yes"]
    parent_rows: list[dict[str, str]] = []
    counters = Counter()
    assigned_parent: list[str] = []

    # First pass: assign IDs by explicit references or fuzzy matching.
    anchors: list[tuple[str, str, str, str]] = []  # master_id, module, title, source_id
    for r in included:
        text = f"{r['Title']} {r['Description / Finding']} {r['Source Bug ID']}"
        module = canonical_module(f"{r['Module']} {text}")
        prefix = module_prefix(module)
        explicit = ""
        m_nb = re.search(r"\bNB-\d{3}\b", text, re.I)
        m_pm = re.search(r"\b(?:Bug|PM)-?\s*(\d{1,3})\b", text, re.I)
        if m_nb:
            explicit = m_nb.group(0).upper()
        elif m_pm:
            explicit = f"PM-{int(m_pm.group(1)):03d}"
        matched = ""
        for mid, mod, title, sid in anchors:
            if explicit and explicit == sid:
                matched = mid
                break
            if mod == module and similarity(r["Title"], title) >= 0.72:
                matched = mid
                break
        if not matched:
            counters[prefix] += 1
            matched = f"BUG-{prefix}-{counters[prefix]:03d}"
            anchors.append((matched, module, r["Title"], explicit or r["Source Bug ID"]))
        r["Master Candidate ID"] = matched
        r["Duplicate Candidate"] = matched
        assigned_parent.append(matched)

    group_counts = Counter(r["Master Candidate ID"] for r in included)
    seen_group = set()
    master_rows = []
    for r in included:
        module = canonical_module(f"{r['Module']} {r['Title']} {r['Description / Finding']}")
        btype = infer_type(f"{r['Title']} {r['Description / Finding']}")
        sev = severity(r["Raw Severity / Priority"], r["Title"], r["Description / Finding"], btype)
        repro = repro_status(r["Raw Status"], r["Source Name"], r["Description / Finding"])
        conf = confidence_for(r["Source Name"], repro, f"{r['Title']} {r['Description / Finding']} {r['Raw Status']}")
        if conf == "Confirmed":
            conf = "Strong Evidence"
        master_id = r["Master Candidate ID"]
        first = master_id not in seen_group
        seen_group.add(master_id)
        if group_counts[master_id] > 1:
            duplicate_status = "Parent Bug" if first else ("Merged Into Parent" if "P0/P1" in r["Source Name"] else "Possible Duplicate")
            duplicate_of = "" if first else master_id
        else:
            duplicate_status = "Unique"
            duplicate_of = ""
        if "duplicate" in r["Triage Outcome"].lower():
            duplicate_status = "Duplicate"
            duplicate_of = master_id
        if "unclear/conflicting" in r["Source Name"]:
            repro = "Code Risk Only"
            conf = "Probable"
        if "Jira export" in r["Source Name"] and repro not in {"Reproduced Twice", "Reproduced Once"}:
            conf = "Unconfirmed"
        action = action_for(sev, btype, duplicate_status, module, r["Source Name"])
        enr = enrichment.get(module, {})
        ev_id = r["Evidence Reference"]
        if r["Source Bug ID"] in evidence_by_id:
            ev_id = f"{r['Evidence Reference']} | {evidence_by_id[r['Source Bug ID']].get('reference_id','')}"
        table_hint = first_nonempty(enr.get("table", [])) or table_guess(module, r)
        field_hint = first_nonempty(enr.get("field", []))
        if "unclear/conflicting" in r["Source Name"]:
            field_hint = db_field_from_unclear(r["Description / Finding"])
            table_hint = db_table_from_unclear(r["Description / Finding"]) or table_hint
        notes = []
        if r["Triage Outcome"]:
            notes.append(f"Triage: {r['Triage Outcome']}.")
        if "Additional evidence document" in r["Source Name"]:
            notes.append("Document-derived issue heading; screenshot/page-level evidence needs human review.")
        if "Jira export" in r["Source Name"]:
            notes.append("Jira export gives issue metadata, not attachment/comment evidence.")
        if conf != "Confirmed":
            notes.append("No explicit Himanshu/Harsh verification evidence found.")
        row = {
            "Master Bug ID": master_id,
            "Source Bug ID": r["Source Bug ID"] or r["Raw Source Row ID"],
            "Source File / Sheet": r["Source File / Sheet"],
            "Source Row / Evidence ID": r["Source Row / Evidence ID"],
            "Module": module,
            "Screen / Route": r["Screen / Route"] or first_nonempty(enr.get("route", [])) or "Needs WD Verification",
            "Bug Title": r["Title"],
            "Bug Description": r["Description / Finding"] or "Needs WD Verification",
            "Bug Type": btype if btype in BUG_TYPES else "Unknown",
            "Severity": sev,
            "Priority Reason": priority_reason(sev, btype, module, r["Title"]),
            "Reproduction Status": repro,
            "Reproduction Steps": reproduction_steps(r),
            "Expected Result": expected_result(r, btype),
            "Actual Result": actual_result(r),
            "Evidence Available": "Yes" if r["Evidence Reference"] or r["Source File / Sheet"] else "No",
            "Evidence Link / Evidence ID": ev_id or r["Source File / Sheet"],
            "Likely Page / Component File": first_nonempty(enr.get("page", [])) or "Needs WD Verification",
            "Likely Hook / Store / Service File": first_nonempty(enr.get("service", [])) or "Needs WD Verification",
            "Likely API Route": first_nonempty(enr.get("api", [])) or "Needs WD Verification",
            "Likely Supabase Table": table_hint or "Needs WD Verification",
            "Likely DB Field": field_hint or "Needs WD Verification",
            "Related Modules": related_modules(module, r, enr),
            "Dashboard Impact": impact_value("dashboard", btype, module, enr),
            "Reminder Impact": impact_value("reminder", btype, module, enr),
            "Document Vault Impact": impact_value("vault", btype, module, enr),
            "Auth / Security Impact": "Yes" if btype in {"Auth", "Configuration"} or module == "Authentication" else "Possible" if "security" in r["Description / Finding"].lower() else "No",
            "Data Loss Risk": data_loss_risk(btype, r),
            "Calculation Risk": "Yes" if btype == "Dashboard Calculation" else "Possible" if any(k in f"{r['Title']} {r['Description / Finding']}".lower() for k in ["amount", "score", "calculation", "income", "expense", "loan", "emi", "roi"]) else "No",
            "Duplicate Status": duplicate_status,
            "Duplicate Of": duplicate_of,
            "Technical Confidence": conf,
            "AI Confidence Reason": confidence_reason(conf, r, repro),
            "Needs WD Verification": "Yes" if conf != "Confirmed" else "No",
            "Suggested Jira Issue Type": "Bug" if r["Triage Outcome"] in {"Valid bug", "Duplicate/supporting source", "Report/fix-plan finding"} else "Task",
            "Suggested Jira Title": jira_title(module, r["Title"]),
            "Recommended Next Action": action,
            "Assigned To": "Pending WD Assignment",
            "Status": "Draft",
            "Notes": " ".join(notes),
        }
        master_rows.append(row)
    return master_rows


def first_nonempty(items) -> str:
    for item in items or []:
        if clean(item):
            return clean(item)
    return ""


def table_guess(module: str, r: dict[str, str]) -> str:
    text = f"{module} {r['Title']} {r['Description / Finding']}".lower()
    guesses = [
        ("family", "family_members"),
        ("education", "education"),
        ("identity", "identity_records"),
        ("credential", "credential_records"),
        ("income", "income_streams"),
        ("expense", "expense_entries; subscriptions"),
        ("subscription", "subscriptions"),
        ("bank", "bank_accounts"),
        ("loan", "loan_accounts"),
        ("insurance", "insurance_policies"),
        ("investment", "investment_accounts"),
        ("property", "property_assets"),
        ("tax", "tax_records"),
        ("reminder", "reminders; notifications"),
        ("document", "document_meta; storage.objects"),
        ("succession", "succession_emergency; succession_nominees"),
        ("will", "will_status; succession_wills"),
        ("user", "users"),
    ]
    return "; ".join(tbl for key, tbl in guesses if key in text) or ""


def db_field_from_unclear(text: str) -> str:
    m = re.search(r"Possible DB:[^;]*;\s*(.+)$", text)
    return clean(m.group(1))[:240] if m else "Needs WD Verification"


def db_table_from_unclear(text: str) -> str:
    m = re.search(r"Possible DB:\s*([^;]+)", text)
    return clean(m.group(1)) if m else ""


def related_modules(module: str, r: dict[str, str], enr: dict) -> str:
    related = []
    for val in enr.get("related", []):
        related.append(val)
    text = f"{r['Title']} {r['Description / Finding']}".lower()
    for name, key in [("Dashboard", "dashboard"), ("Reminders", "reminder"), ("Document Vault", "document"), ("Authentication", "auth")]:
        if key in text and name not in related:
            related.append(name)
    return clean("; ".join(related)) or "Needs WD Verification"


def impact_value(kind: str, btype: str, module: str, enr: dict) -> str:
    if kind == "dashboard" and (btype == "Dashboard Calculation" or module == "Rajya / Dashboard"):
        return "Yes"
    if kind == "reminder" and btype == "Reminder":
        return "Yes"
    if kind == "vault" and btype == "Document Upload":
        return "Yes"
    value = first_nonempty(enr.get(kind, []))
    if value:
        if value.lower().startswith("high"):
            return "Possible"
        if value.lower().startswith("medium"):
            return "Possible"
    return "No"


def data_loss_risk(btype: str, r: dict[str, str]) -> str:
    text = f"{r['Title']} {r['Description / Finding']}".lower()
    if any(k in text for k in ["lost", "delete", "orphan", "cannot remove", "wipe", "data loss", "not persisted"]):
        return "High"
    if btype in {"Data Persistence", "Delete", "Document Upload"}:
        return "Possible"
    return "Low"


def action_for(sev: str, btype: str, dup: str, module: str, source: str) -> str:
    if dup == "Duplicate":
        return "No Action — Duplicate"
    if btype == "Auth":
        return "Needs Auth Review"
    if "Supabase" in module or btype == "Configuration":
        return "Needs Supabase Verification"
    if sev in {"P0", "P1"}:
        return "Move to P0/P1 Impact Report"
    if "DA unclear" in source:
        return "Needs WD Reproduction"
    if "Jira export" in source:
        return "Merge with Existing Jira"
    if "Additional evidence" in source:
        return "Needs WD Reproduction"
    return "Create Jira Bug"


def reproduction_steps(r: dict[str, str]) -> str:
    if "DA bug verification" in r["Source Name"]:
        return r["Title"]
    if "Jira export" in r["Source Name"]:
        return "Use Jira issue description and reproduce with dummy account in isolated environment."
    return "See source row/evidence; reproduce with dummy data only in isolated environment."


def expected_result(r: dict[str, str], btype: str) -> str:
    text = r["Description / Finding"]
    if "Expected" in text:
        return text[:350]
    defaults = {
        "Auth": "User should authenticate, recover account, or complete auth callback through the intended route without unsafe session behavior.",
        "Save/Create": "Record should save once to the authoritative persistence layer and be visible after refresh/re-login.",
        "Edit/Update": "Edits should persist and update dependent modules consistently.",
        "Delete": "Deletion should remove or safely unlink dependent records according to product rules.",
        "Reminder": "Relevant reminder should be created, updated, or removed consistently.",
        "Document Upload": "Document should upload, link, display, and delete through the intended vault/storage flow.",
        "Validation": "Invalid input should be blocked with clear field-level feedback.",
        "Responsive/UI": "UI should remain clear, usable, and accessible on target viewports.",
    }
    return defaults.get(btype, "Expected behavior requires PM/WD confirmation.")


def actual_result(r: dict[str, str]) -> str:
    return r["Description / Finding"][:700] or "Actual result needs WD verification."


def confidence_reason(conf: str, r: dict[str, str], repro: str) -> str:
    if conf == "Strong Evidence":
        return f"Source evidence reports {repro}; technical/root-cause verification remains pending."
    if conf == "Probable":
        return "Static mapping or pattern suggests a risk, but reproduction or direct root-cause proof is incomplete."
    return "Source row/report exists, but current live reproduction, code proof, or WD verification is missing."


def jira_title(module: str, title: str) -> str:
    short = re.sub(r"\s+", " ", clean(title))
    return f"[{module}] {short}"[:180]


def create_source_inventory(source_stats, master_rows):
    included_by_source = Counter(r["Source File / Sheet"] for r in master_rows)
    rows = []
    for name, (path, owner, reviewed, raw_count, included, notes) in source_stats.items():
        evidence = "Yes" if Path(path).exists() else "No"
        rows.append({
            "Source Name": name,
            "File Name": simple_path(Path(path)),
            "Owner Team": owner,
            "Date": "Needs WD Verification",
            "Type of Bugs": notes,
            "Evidence Available": evidence,
            "Rows Reviewed": str(reviewed),
            "Rows Sent To Raw Merge": str(raw_count),
            "Rows Included In Master": str(included),
            "Notes": "Offline local file; no external evidence URLs opened.",
        })
    return rows


def write_gap_report(master_rows, source_inventory):
    by_conf = Counter(r["Technical Confidence"] for r in master_rows)
    by_repro = Counter(r["Reproduction Status"] for r in master_rows)
    by_sev = Counter(r["Severity"] for r in master_rows)
    by_dup = Counter(r["Duplicate Status"] for r in master_rows)
    needs_wd = [r for r in master_rows if r["Needs WD Verification"] == "Yes"]
    p1 = [r for r in master_rows if r["Severity"] in {"P0", "P1"}]
    da_unclear = [r for r in raw if r["Source Name"] == "DA unclear/conflicting fields"]
    da_nonpass = [r for r in raw if r["Source Name"] == "DA bug verification journeys" and r["Include In Master"] == "Yes"]
    missing = [
        "TECHNICAL_BUG_REPRODUCTION.csv is still not available in audit-input/03-bug-and-evidence.",
        "PM_New_Bug_List.csv and PM_Duplicate_Bug_List.csv are not available as separate files; Jira and NB/PM files were used instead.",
        "DA_User_Journey_Test_Results.csv is not available as a separate file; DA_BUG_VERIFICATION.csv was used.",
        "Sanitised browser console logs and sanitised build logs are still missing from the live/deployment context folder.",
        "Jira comments, attachments, and linked-issue duplicate metadata were not present in the CSV exports.",
    ]
    lines = [
        "# Bug Cross-Reference Gaps",
        "",
        "Date: 2026-07-20",
        "",
        "Status: Draft - Pending Evidence Cleanup and WD Verification",
        "",
        "This file records the remaining evidence gaps after rebuilding the master bug register. No external service was contacted and no application source files were changed.",
        "",
        "## Rebuild Summary",
        "",
        f"- Master rows: {len(master_rows)}",
        f"- Raw merge rows: {len(raw)}",
        f"- Source inventory rows: {len(source_inventory)}",
        f"- P0/P1 rows in master: {len(p1)}",
        f"- Rows needing WD verification: {len(needs_wd)}",
        f"- DA non-pass journey rows triaged: {len(da_nonpass)}",
        f"- DA unclear/conflicting field rows triaged: {len(da_unclear)}",
        "",
        "## Counts",
        "",
        f"- Severity: {dict(by_sev)}",
        f"- Reproduction status: {dict(by_repro)}",
        f"- Duplicate status: {dict(by_dup)}",
        f"- Technical confidence: {dict(by_conf)}",
        "",
        "## Missing Or Partial Inputs",
        "",
    ]
    lines.extend(f"- {m}" for m in missing)
    lines.extend([
        "",
        "## Duplicate And Grouping Gaps",
        "",
        "- Duplicate grouping was performed from explicit PM/NB references, P0/P1 references, Jira title similarity, and fuzzy title matching by module.",
        "- Jira parent/subtask links are present in the raw source, but Jira linked issues and duplicate relationships were not exported.",
        "- Additional PDF/DOCX reports contain historical bugs and UX issues that may duplicate existing PM/NB/Jira rows; document-derived rows are marked Needs WD Verification.",
        "",
        "## WD Verification Required",
        "",
    ])
    for r in needs_wd[:80]:
        lines.append(f"- {r['Master Bug ID']} / {r['Source Bug ID']}: {r['Bug Title']} ({r['Technical Confidence']}, {r['Recommended Next Action']})")
    if len(needs_wd) > 80:
        lines.append(f"- Additional rows needing WD verification not listed here: {len(needs_wd) - 80}")
    lines.extend([
        "",
        "## PM/DA Clarification Required",
        "",
        "- Decide which UX-report observations should become Jira bugs versus product backlog improvements.",
        "- Confirm whether module dependency observations from the Money App dependency report override or supplement the current MODULE_DEPENDENCY_MATRIX.csv.",
        "- Confirm canonical severity for P1 candidates where Jira priority and PM severity differ.",
        "",
        "## Files To Update After WD Review",
        "",
        "- `CONSOLIDATED_BUG_CROSS_REFERENCE.csv`",
        "- `P0_P1_BUG_IMPACT_REPORT.md`",
        "- `MANUAL_REGRESSION_TEST_CASES.csv`",
        "- `AUTOMATED_TEST_RECOMMENDATIONS.md`",
        "- `06-final-handover/ENGINEER_HANDOVER_INDEX.md`",
        "- `06-final-handover/FINAL_AI_AUDIT_PACK_STATUS.md`",
    ])
    (OUT / "BUG_CROSS_REFERENCE_GAPS.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_p0p1(master_rows):
    p_rows = [r for r in master_rows if r["Severity"] in {"P0", "P1"}]
    by_id = {}
    for r in p_rows:
        by_id.setdefault(r["Master Bug ID"], r)
    lines = [
        "# P0/P1 Bug Impact Report",
        "",
        "**Rebuilt from master register:** 2026-07-20",
        "",
        "This report is reconciled to `CONSOLIDATED_BUG_CROSS_REFERENCE.csv` master bug IDs. It remains audit-only: no external service was contacted, no production test was run, and no application source file was modified.",
        "",
        "## Severity Position",
        "",
        f"- Confirmed P0 rows: {sum(1 for r in p_rows if r['Severity'] == 'P0' and r['Technical Confidence'] == 'Confirmed')}",
        f"- P0/P1 master rows: {len(p_rows)}",
        f"- Unique P0/P1 master IDs: {len(by_id)}",
        "- No item is marked Verified or Confirmed by Himanshu/Harsh in this audit output.",
        "",
        "## Reconciled P0/P1 Master IDs",
        "",
        "| Master Bug ID | Severity | Module | Title | User journey / downstream impact | Evidence status | Recommended next action |",
        "|---|---|---|---|---|---|---|",
    ]
    for mid in sorted(by_id):
        r = by_id[mid]
        impact = "; ".join([
            f"Dashboard: {r['Dashboard Impact']}",
            f"Reminder: {r['Reminder Impact']}",
            f"Vault: {r['Document Vault Impact']}",
            f"Auth/Security: {r['Auth / Security Impact']}",
            f"Data loss: {r['Data Loss Risk']}",
            f"Calculation: {r['Calculation Risk']}",
        ])
        lines.append(f"| `{mid}` | {r['Severity']} | {r['Module']} | {r['Bug Title']} | {impact} | {r['Technical Confidence']} / WD={r['Needs WD Verification']} | {r['Recommended Next Action']} |")
    lines.extend([
        "",
        "## Required Investigation Order",
        "",
        "1. Security/auth/access risks: authentication, OAuth, password reset, OTP, profile/session, and environment-isolation rows.",
        "2. Data loss risks: document vault/storage, deletion/unlinking, local-only persistence, and duplicate/orphan rows.",
        "3. Database integrity risks: schema mismatch, field mapping, soft links, and missing constraints.",
        "4. Shared dependency risks: dashboard, reminder, notification, vault, and module route coupling.",
        "5. Blocked user journeys: P1 rows with route, save, upload, login, or schema blockers.",
        "6. Lower-priority UX issues: P3 and UX-derived rows after functional blockers are assigned.",
        "",
        "## Limitations",
        "",
        "- `TECHNICAL_BUG_REPRODUCTION.csv` is still missing.",
        "- Jira comments, attachments, linked issues, and duplicate relationships were not included in the CSV exports.",
        "- PDF/DOCX screenshots were not visually redacted or independently re-tested.",
        "- All rows remain Draft / Needs WD Verification until authorized verification evidence is added.",
    ])
    (OUT / "P0_P1_BUG_IMPACT_REPORT.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def update_handover(master_rows, source_inventory):
    sev = Counter(r["Severity"] for r in master_rows)
    conf = Counter(r["Technical Confidence"] for r in master_rows)
    p1_ids = len({r["Master Bug ID"] for r in master_rows if r["Severity"] in {"P0", "P1"}})
    addendum = f"""

## 2026-07-20 Bug Register Rebuild Addendum

- `audit-output/04-bugs-and-qa/CONSOLIDATED_BUG_CROSS_REFERENCE.csv` was rebuilt to the required 40-column acceptance schema.
- Previous 24-column draft preserved as `audit-output/04-bugs-and-qa/CONSOLIDATED_BUG_CROSS_REFERENCE_2026-07-20_PRE_40_COLUMN_REBUILD_DRAFT.csv`.
- New support files: `BUG_SOURCE_INVENTORY.csv`, `RAW_BUG_MERGE.csv`, and `BUG_CROSS_REFERENCE_GAPS.md`.
- Rebuilt master rows: {len(master_rows)}; source inventory rows: {len(source_inventory)}; unique P0/P1 master IDs: {p1_ids}.
- Severity counts: {dict(sev)}.
- Technical confidence counts: {dict(conf)}.
- Status remains Needs WD Verification. No row was marked `Confirmed` because explicit Himanshu/Harsh/WD verification evidence was not found.
"""
    for path in [
        HANDOVER / "ENGINEER_HANDOVER_INDEX.md",
        HANDOVER / "FINAL_AI_AUDIT_PACK_STATUS.md",
        HANDOVER / "EXECUTIVE_TECHNICAL_SUMMARY.md",
    ]:
        txt = path.read_text(encoding="utf-8") if path.exists() else ""
        marker = "## 2026-07-20 Bug Register Rebuild Addendum"
        if marker in txt:
            txt = txt.split(marker)[0].rstrip() + "\n"
        path.write_text(txt.rstrip() + "\n" + addendum, encoding="utf-8")

    seq = HANDOVER / "RECOMMENDED_FIX_SEQUENCE.md"
    txt = seq.read_text(encoding="utf-8") if seq.exists() else "# Recommended Fix Sequence\n"
    marker = "## 2026-07-20 Master Bug ID Reconciliation"
    sequence_addendum = f"""

{marker}

Use the rebuilt `CONSOLIDATED_BUG_CROSS_REFERENCE.csv` as the source of truth for bug IDs before creating or merging Jira tickets. P0/P1 rows have been reconciled into `P0_P1_BUG_IMPACT_REPORT.md`; all remain Draft / Needs WD Verification.
"""
    if marker in txt:
        txt = txt.split(marker)[0].rstrip() + "\n"
    seq.write_text(txt.rstrip() + "\n" + sequence_addendum, encoding="utf-8")


def validate(master_rows):
    errors = []
    required_nonblank = {
        "Master Bug ID",
        "Source Bug ID",
        "Source File / Sheet",
        "Source Row / Evidence ID",
        "Module",
        "Screen / Route",
        "Bug Title",
        "Bug Description",
        "Bug Type",
        "Severity",
        "Priority Reason",
        "Reproduction Status",
        "Expected Result",
        "Actual Result",
        "Evidence Available",
        "Evidence Link / Evidence ID",
        "Technical Confidence",
        "AI Confidence Reason",
        "Needs WD Verification",
        "Suggested Jira Issue Type",
        "Suggested Jira Title",
        "Recommended Next Action",
        "Assigned To",
        "Status",
    }
    for idx, row in enumerate(master_rows, start=2):
        if set(row.keys()) != set(MASTER_HEADERS):
            missing = set(MASTER_HEADERS) - set(row.keys())
            extra = set(row.keys()) - set(MASTER_HEADERS)
            errors.append(f"row {idx} header mismatch missing={missing} extra={extra}")
        for col in required_nonblank:
            if clean(row.get(col, "")) == "":
                errors.append(f"row {idx} blank {col}")
        if row["Bug Type"] not in BUG_TYPES:
            errors.append(f"row {idx} invalid Bug Type {row['Bug Type']}")
        if row["Severity"] not in {"P0", "P1", "P2", "P3"}:
            errors.append(f"row {idx} invalid Severity {row['Severity']}")
        if row["Reproduction Status"] not in REPRO_STATUSES:
            errors.append(f"row {idx} invalid Reproduction Status {row['Reproduction Status']}")
        if row["Duplicate Status"] not in DUP_STATUSES:
            errors.append(f"row {idx} invalid Duplicate Status {row['Duplicate Status']}")
        if row["Technical Confidence"] not in CONFIDENCE:
            errors.append(f"row {idx} invalid Technical Confidence {row['Technical Confidence']}")
        if row["Recommended Next Action"] not in ACTIONS:
            errors.append(f"row {idx} invalid Recommended Next Action {row['Recommended Next Action']}")
        if row["Needs WD Verification"] not in {"Yes", "No"}:
            errors.append(f"row {idx} invalid Needs WD Verification {row['Needs WD Verification']}")
        if row["Technical Confidence"] == "Confirmed":
            errors.append(f"row {idx} incorrectly Confirmed without authorized verification")
        if row["Duplicate Status"] in {"Duplicate", "Possible Duplicate", "Merged Into Parent"} and clean(row.get("Duplicate Of", "")) == "":
            errors.append(f"row {idx} duplicate row missing Duplicate Of")
    text = "\n".join(",".join(row.values()) for row in master_rows)
    if re.search(r"[A-Za-z0-9._%+-]+@(?!example\.com)[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text):
        errors.append("non-dummy email found in master")
    if re.search(r"eyJ[A-Za-z0-9_-]{20,}\.[A-Za-z0-9_-]{20,}\.[A-Za-z0-9_-]{20,}", text):
        errors.append("jwt-like token found in master")
    return errors


def main():
    source_stats = gather_sources()
    master_rows = make_master()
    source_inventory = create_source_inventory(source_stats, master_rows)
    write_csv(OUT / "RAW_BUG_MERGE.csv", RAW_HEADERS, raw)
    write_csv(OUT / "BUG_SOURCE_INVENTORY.csv", SOURCE_HEADERS, source_inventory)
    write_csv(OUT / "CONSOLIDATED_BUG_CROSS_REFERENCE.csv", MASTER_HEADERS, master_rows)
    write_gap_report(master_rows, source_inventory)
    write_p0p1(master_rows)
    update_handover(master_rows, source_inventory)
    errors = validate(master_rows)
    validation_lines = [
        "# Bug Register Rebuild Validation",
        "",
        "Date: 2026-07-20",
        "",
        f"Master rows: {len(master_rows)}",
        f"Raw rows: {len(raw)}",
        f"Source inventory rows: {len(source_inventory)}",
        f"Validation errors: {len(errors)}",
        "",
    ]
    validation_lines.extend(f"- {e}" for e in errors[:200])
    if len(errors) > 200:
        validation_lines.append(f"- Additional errors omitted: {len(errors)-200}")
    (OUT / "BUG_REGISTER_REBUILD_VALIDATION.md").write_text("\n".join(validation_lines) + "\n", encoding="utf-8")
    print(f"master_rows={len(master_rows)} raw_rows={len(raw)} source_rows={len(source_inventory)} errors={len(errors)}")


if __name__ == "__main__":
    main()
